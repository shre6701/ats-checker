"""Backend tests for ATS Resume Scanner."""
import os
import io
import pytest
import requests
from docx import Document

BASE_URL = os.environ.get("REACT_APP_BACKEND_URL", "https://job-match-scanner-2.preview.emergentagent.com").rstrip("/")
SESSION_TOKEN = "test_session_1737000000"
HEADERS = {"Authorization": f"Bearer {SESSION_TOKEN}"}

SAMPLE_RESUME = (
    "Jane Doe\nSoftware Engineer with 5 years of experience in Python, FastAPI, React, AWS, Docker. "
    "Built scalable APIs.\nExperience: Senior SWE at Acme (2021-2025) - led team of 4, shipped microservices."
)
SAMPLE_JD = (
    "We are hiring a Senior Backend Engineer with Python, FastAPI, AWS, Kubernetes, microservices, "
    "CI/CD experience. Must have 5+ yrs and leadership."
)


# Health
def test_health():
    r = requests.get(f"{BASE_URL}/api/")
    assert r.status_code == 200
    assert r.json().get("ok") is True


# Auth: unauthenticated
def test_auth_me_unauthenticated():
    r = requests.get(f"{BASE_URL}/api/auth/me")
    assert r.status_code == 401


# Auth: authenticated
def test_auth_me_authenticated():
    r = requests.get(f"{BASE_URL}/api/auth/me", headers=HEADERS)
    assert r.status_code == 200
    data = r.json()
    assert data["email"] == "test.user@example.com"
    assert data["name"] == "Jane Doe"


# Resume parse: DOCX -> returns source_format=docx
def test_parse_docx():
    doc = Document()
    doc.add_paragraph("Jane Doe - Software Engineer")
    doc.add_paragraph("Python, FastAPI, AWS")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    files = {"file": ("resume.docx", buf.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    r = requests.post(f"{BASE_URL}/api/resumes/parse", files=files, headers=HEADERS)
    assert r.status_code == 200, r.text
    data = r.json()
    assert "Jane Doe" in data["text"]
    assert data.get("source_format") == "docx"


# Resume parse: TXT -> returns source_format=txt
def test_parse_txt():
    files = {"file": ("resume.txt", b"Jane Doe\nPython developer", "text/plain")}
    r = requests.post(f"{BASE_URL}/api/resumes/parse", files=files, headers=HEADERS)
    assert r.status_code == 200, r.text
    data = r.json()
    assert data.get("source_format") == "txt"
    assert "Jane Doe" in data["text"]


# Scan pipeline (slow - 3 LLM calls) - DOCX source
@pytest.fixture(scope="module")
def scan_result():
    r = requests.post(
        f"{BASE_URL}/api/resumes/scan",
        json={
            "resume_text": SAMPLE_RESUME,
            "job_description": SAMPLE_JD,
            "job_title": "Senior Backend Engineer",
            "company": "TestCo",
            "source_format": "docx",
        },
        headers=HEADERS,
        timeout=180,
    )
    assert r.status_code == 200, r.text
    return r.json()


def test_scan_stores_source_format(scan_result):
    assert scan_result.get("source_format") == "docx"


def test_scan_response_shape(scan_result):
    assert "scan_id" in scan_result
    assert "analysis" in scan_result
    ats = scan_result["analysis"].get("ats_score")
    assert isinstance(ats, int) and 0 <= ats <= 100
    assert isinstance(scan_result["optimized_resume"], str) and len(scan_result["optimized_resume"]) > 50
    assert isinstance(scan_result["cover_letter"], str) and len(scan_result["cover_letter"]) > 50
    assert isinstance(scan_result["new_ats_score"], int)


def test_scan_new_score_high(scan_result):
    assert scan_result["new_ats_score"] >= 85, f"new_ats_score={scan_result['new_ats_score']}"


# History
def test_history_list(scan_result):
    r = requests.get(f"{BASE_URL}/api/history", headers=HEADERS)
    assert r.status_code == 200
    ids = [it["scan_id"] for it in r.json()]
    assert scan_result["scan_id"] in ids


def test_history_get(scan_result):
    r = requests.get(f"{BASE_URL}/api/history/{scan_result['scan_id']}", headers=HEADERS)
    assert r.status_code == 200
    assert r.json()["scan_id"] == scan_result["scan_id"]


# Downloads - default (no fmt) follows source_format=docx
def test_download_resume_default_uses_source_format(scan_result):
    r = requests.get(f"{BASE_URL}/api/history/{scan_result['scan_id']}/download/resume", headers=HEADERS)
    assert r.status_code == 200
    assert r.headers["content-type"].startswith("application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    assert r.content[:4] == b"PK\x03\x04"
    # Verify openable by python-docx
    d = Document(io.BytesIO(r.content))
    assert any(p.text for p in d.paragraphs)


# Force PDF via fmt query param
def test_download_resume_fmt_pdf(scan_result):
    r = requests.get(f"{BASE_URL}/api/history/{scan_result['scan_id']}/download/resume?fmt=pdf", headers=HEADERS)
    assert r.status_code == 200
    assert r.headers["content-type"].startswith("application/pdf")
    assert r.content.startswith(b"%PDF")


# Force DOCX via fmt query param
def test_download_resume_fmt_docx(scan_result):
    r = requests.get(f"{BASE_URL}/api/history/{scan_result['scan_id']}/download/resume?fmt=docx", headers=HEADERS)
    assert r.status_code == 200
    assert r.headers["content-type"].startswith("application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    assert r.content[:4] == b"PK\x03\x04"
    d = Document(io.BytesIO(r.content))
    assert any(p.text for p in d.paragraphs)


# Force TXT via fmt query param on cover letter
def test_download_cover_fmt_txt(scan_result):
    r = requests.get(f"{BASE_URL}/api/history/{scan_result['scan_id']}/download/cover?fmt=txt", headers=HEADERS)
    assert r.status_code == 200
    assert r.headers["content-type"].startswith("text/plain")
    body = r.content.decode("utf-8")
    assert len(body) > 50
    assert "Sincerely" in body or "Dear" in body


# Cover default uses source_format=docx
def test_download_cover_default_docx(scan_result):
    r = requests.get(f"{BASE_URL}/api/history/{scan_result['scan_id']}/download/cover", headers=HEADERS)
    assert r.status_code == 200
    assert r.headers["content-type"].startswith("application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    assert r.content[:4] == b"PK\x03\x04"


# Delete
def test_delete_scan(scan_result):
    r = requests.delete(f"{BASE_URL}/api/history/{scan_result['scan_id']}", headers=HEADERS)
    assert r.status_code == 200
    r2 = requests.get(f"{BASE_URL}/api/history/{scan_result['scan_id']}", headers=HEADERS)
    assert r2.status_code == 404
